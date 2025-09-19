# @version : 1.0
# @Author  : 混江龙
# @File    : word_for_python.py
# @Time    : 2025/9/19 14:34
# pip install python-docx  处理文档的工具包
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING
"""
doc = Document('test.docx')
# 读取所有的段落
all_paragraphs = doc.paragraphs
# 段落的最小单位叫 run
for paragraph in all_paragraphs: #循环所有的段落
    for run in paragraph.runs: #循环每个段落中所有的run
        print(run.text) #每个run中的内容
# 读取word中所有的表格
all_table = doc.tables
# 表格的最小单位叫 cell
for table in all_table:
    for row in table.rows: #读取每个表格中所有的行
        for cell in row.cells: #读取每行中的每一格（cell）
            print(cell.text)
# 保存word文件
# doc.save('test.docx')
"""
# 实例化word对象
document = Document()
# 添加标题 H2级标题
document.add_heading("标题",2)
# 添加段落
document.add_paragraph("添加段落")
# 创建表格 3行，4列
table = document.add_table(3,4,style="Table Grid")
# 第一行
table.rows[0].cells[0].text = "1"
table.rows[0].cells[1].text = "2020/1/1"
table.rows[0].cells[2].text = "一一俱乐部"
table.rows[0].cells[3].text = "Xx子（2025）111号"
# 第二行
table.rows[1].cells[1].text = "关于做好卫生检查的通知"
# 第三行
table.rows[2].cells[1].text = "签收"
# 合并单元格
# 第0行的0列合并到第2行的0列
table.cell(0,0).merge(table.cell(2,0))
# 第1行的1列合并到第1行的3列
table.cell(1,1).merge(table.cell(1,3))
# 居中文本
# 水平居中 WD_PARAGRAPH_ALIGNMENT.CENTER
# 行间距 WD_LINE_SPACING
table.cell(1,1).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# table.cell(2,1).merge(table.cell(2,3))
# 添加分页符，一页一张表
document.add_page_break()
document.save('test.docx')