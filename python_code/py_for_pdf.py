# @version : 1.0
# @Author  : 混江龙
# @File    : py_for_pdf.py
# @Time    : 2025/9/19 10:19
# pdfplumber 用于识别pdf中的文字和表格
import pdfplumber
import docx
import pandas as pd
# PyPDF2可以用于pdf的读写操作
from PyPDF2 import PdfReader,PdfWriter

# 读取pdf
p = pdfplumber.open(r"C:\Users\Administrator\Desktop\网站源码转出协议.pdf")
# print(p)
# pdf中的所有的页
# print(p.pages)
# 取出第一页
page = p.pages[0]
# 读取页中的文字
txt = page.extract_text()
doc = docx.Document()
# 把文本根据换行符划分成段落
dl = txt.split("\n")
for i in range(len(dl)):
    # 把第一行设置成标题
    if i == 0:
        doc.add_heading(dl[i], level=1)
    else:
        # 其他行添加到段落
        doc.add_paragraph(dl[i])
# 保存成word文档
# doc.save(rf"C:\Users\Administrator\Desktop\{dl[0]}.docx")
# 读取页中的表格
table = page.extract_tables()
# print(txt,table)
# 把table中的表数据转到pandas数据，生成表格
pd_table = pd.DataFrame(table[0])
# 去掉默认的第一行方法：
tb_header = pd_table.loc[0,:] #[0,:] 获取表头 第0 行的所有列
pd_table = pd_table.loc[1:,] # 表格内容 赋给原变量，从第一行到所有行，所有列
pd_table.columns = tb_header #把表头设置给新的表内容
# print(pd_table)
# 把表格保存成xlsx文件，并去掉默认第一列，就是index索引
# pd_table.to_excel(r"C:\Users\Administrator\Desktop\pdf.xlsx",index=False)

# 读取pdf
input_pdf = PdfReader(r'C:\Users\Administrator\Desktop\黄冈师范校友会售后维护合同.pdf')
# print(len(input_pdf.pages)) #读取pdf有多少页
# print(input_pdf.pages[0]) #读取第一页的内容
pdf_writer = PdfWriter() #创建一个新的空pdf文档
pdf_writer.add_page(input_pdf.pages[0]) #把第一页的内容加入空文档
# 保存新的文件
with open(r'C:\Users\Administrator\Desktop\第一页.pdf','wb') as f:
    pdf_writer.write(f)


# 合并pdf文档